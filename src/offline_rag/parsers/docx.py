from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_docx(path: Path) -> tuple[str, dict]:
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed; cannot parse DOCX")
        return "", {"file_type": "docx", "lang": "docx"}

    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs), {"file_type": "text", "lang": "docx"}
    except Exception as e:
        logger.warning(f"DOCX parse error {path}: {e}")
        return "", {"file_type": "text", "lang": "docx"}
